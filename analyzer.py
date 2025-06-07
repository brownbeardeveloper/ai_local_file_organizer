"""
File Analyzer Module
Orchestrates file content analysis using specialized AI model analyzers
"""

from pathlib import Path
from typing import Dict, Optional, Any
from datetime import datetime
import exifread
from yolo_analyzer import YOLOAnalyzer
from ollama_analyzer import OllamaAnalyzer
from ocr_analyzer import OCRAnalyzer
import torch
import PyPDF2
from docx import Document

# Configure PyTorch for Mac optimization
if torch.backends.mps.is_available():
    torch.backends.mps.is_built()
    print("MPS acceleration available")
else:
    print("MPS not available, using CPU")


class FileAnalyzer:
    """
    Analyzes files using specialized AI analyzers for images, documents, audio, and video.

    This class orchestrates the use of YOLO (object detection), Ollama (text/content analysis), and OCR (optical character recognition) analyzers
    to extract metadata, classify content, and provide AI-driven insights for a wide range of file types.

    Features:
        - Modular analysis for images, documents, audio, and video files
        - Integrates multiple AI models for deep content understanding
        - Handles large files, empty files, and various edge cases
        - Returns structured metadata and AI insights for downstream use
    """

    def __init__(
        self,
        large_file_threshold_mb: int = 100,  # in MB
        yolo_analyzer: Optional[YOLOAnalyzer] = None,
        ollama_analyzer: Optional[OllamaAnalyzer] = None,
        ocr_analyzer: Optional[OCRAnalyzer] = None,
    ):
        """
        Initialize the analyzer with optional AI components.

        Args:
            large_file_threshold_mb: File size threshold in MB (1-300). Files larger than this will be treated differently.
            yolo_analyzer: Optional YOLO object detection analyzer.
            ollama_analyzer: Optional language model analyzer.
            ocr_analyzer: Optional OCR analyzer.

        Raises:
            ValueError: If large_file_threshold_mb is not in the range 1-300.
        """
        if not (1 <= large_file_threshold_mb <= 300):
            raise ValueError("large_file_threshold_mb must be between 1 and 300 MB")

        self.large_file_threshold = (
            large_file_threshold_mb * 1024 * 1024
        )  # store internally in bytes
        self.yolo_analyzer = yolo_analyzer or YOLOAnalyzer()
        self.ollama_analyzer = ollama_analyzer or OllamaAnalyzer()
        self.ocr_analyzer = ocr_analyzer or OCRAnalyzer()

    def analyze(self, file_info: Dict) -> Dict[str, Any]:
        """Analyze file content and return insights about what's inside.

        Args:
            file_info: Dictionary containing file metadata with keys: path, name, suffix,
                      category, size, modified, mime_type.

        Returns:
            Dict containing original file metadata plus analysis results and AI insights.

        Raises:
            TypeError: If file_info["path"] is not a string or path-like object.
            ValueError: If the path string contains null bytes.
            FileNotFoundError: If image, document, or PDF files don't exist.
            PermissionError: If no read permissions for the file.
            OSError: General I/O errors (disk issues, network problems).
            RuntimeError: If YOLO model unavailable, Ollama service not running,
                         encrypted/empty PDFs, or OCR extraction failures.
            UnicodeDecodeError: If text files have encoding issues.
            MemoryError: If large files exceed system memory on constrained systems.
            ConnectionError: If network issues occur with Ollama service.
            TimeoutError: If Ollama requests timeout.
        """
        file_path = Path(file_info["path"])

        # Start with file metadata from scanner and add analysis results
        analysis = {
            # Preserve original file metadata
            "path": file_info["path"],
            "name": file_info["name"],
            "suffix": file_info["suffix"],
            "category": file_info["category"],
            "size": file_info["size"],
            "modified": file_info["modified"],
            "mime_type": file_info["mime_type"],
            # Analysis results
            "content_analysis": "pending",
            "ai_insights": {},
        }

        # Content analysis based on file category from scanner
        file_category = file_info.get("category", "other")

        if file_category == "image":
            analysis.update(self._analyze_image_content(file_info["path"]))
        elif file_category == "document":
            analysis.update(self._analyze_document_content(file_path, file_info))
        elif file_category == "video":
            analysis.update(self._analyze_video_content(file_path))
        elif file_category == "audio":
            analysis.update(self._analyze_audio_content(file_path))
        else:
            analysis["content_analysis"] = "unsupported_type"

        return analysis

    def _analyze_image_content(self, path: str) -> Dict[str, Any]:
        """Analyze image content using YOLO object detection and EXIF metadata extraction.

        Args:
            path: String path to the image file.

        Returns:
            Dict containing analysis results and AI insights.

        Raises:
            RuntimeError: If YOLO analyzer is not available for image object detection.
        """
        result = {
            "content_analysis": "image_analyzed",
            "ai_insights": {},
        }

        try:
            # YOLO analysis for object detection
            if not self.yolo_analyzer:
                raise RuntimeError(
                    "YOLO analyzer not available for image object detection"
                )

            detected_objects = self.yolo_analyzer.detect_objects(path)
            if detected_objects:
                result["ai_insights"].update(detected_objects)
            else:
                result["ai_insights"]["primary_object"] = "no_objects_detected"

            # Extract EXIF for photo metadata
            self._extract_photo_context(path, result)

        except Exception as e:
            result["content_analysis"] = f"image_error: {str(e)}"

        return result

    def _analyze_document_content(self, path: Path, file_info: Dict) -> Dict[str, Any]:
        """Analyze what's actually IN the document"""
        result = {
            "content_analysis": "document_analyzed",
            "ai_insights": {},
        }

        try:
            suffix = file_info.get("suffix", "").lower()
            file_size = file_info.get("size", 0)

            # Skip empty files
            if file_size == 0:
                result["ai_insights"]["content"] = "empty_file"
                return result

            # Skip very large files for content analysis
            if file_size > self.large_file_threshold:
                result["ai_insights"]["content"] = "large_file_skipped"
                return result

            if suffix == ".pdf":
                content_info = self._analyze_pdf_content(path)
            elif suffix in [".txt", ".md"]:
                content_info = self._analyze_text_content(path)
            elif suffix in [".doc", ".docx"]:
                content_info = self._analyze_word_content(path)
            elif suffix in [".csv"]:
                content_info = self._analyze_csv_content(path)
            elif suffix in [".py", ".js", ".sh", ".html", ".css"]:
                content_info = self._analyze_code_content(path)
            else:
                content_info = {"content": "unsupported_document_type"}

            result["ai_insights"].update(content_info)

        except Exception as e:
            result["content_analysis"] = f"document_error: {str(e)}"

        return result

    def _analyze_video_content(self, path: Path) -> Dict[str, Any]:
        """Analyze video file characteristics"""
        return {
            "content_analysis": "video_detected",
            "ai_insights": {"content": "video_analysis_not_implemented"},
        }

    def _analyze_audio_content(self, path: Path) -> Dict[str, Any]:
        """Analyze audio file characteristics"""
        return {
            "content_analysis": "audio_detected",
            "ai_insights": {"content": "audio_analysis_not_implemented"},
        }

    def _extract_photo_context(self, path: str, result: Dict):
        """Extract photo metadata for context"""
        try:
            with open(path, "rb") as f:
                tags = exifread.process_file(f, stop_tag="EXIF DateTimeOriginal")
                if "EXIF DateTimeOriginal" in tags:
                    date_str = str(tags["EXIF DateTimeOriginal"])
                    try:
                        photo_date = datetime.strptime(date_str, "%Y:%m:%d %H:%M:%S")
                        result["ai_insights"]["photo_date"] = photo_date.strftime(
                            "%Y-%m-%d"
                        )

                        # Determine if it's recent or old
                        days_old = (datetime.now() - photo_date).days
                        if days_old < 30:
                            result["ai_insights"]["age"] = "recent"
                        elif days_old < 365:
                            result["ai_insights"]["age"] = "this_year"
                        else:
                            result["ai_insights"]["age"] = "old"
                    except ValueError:
                        pass
        except Exception:
            pass

    def _extract_pdf_text(self, path: Path) -> str:
        """Extract text from all pages of a PDF"""
        with open(path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)

            if pdf_reader.is_encrypted:
                raise RuntimeError("Cannot extract text from encrypted PDF")

            if len(pdf_reader.pages) == 0:
                raise RuntimeError("PDF file contains no pages")

            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text()

            return text_content

    def _analyze_pdf_content(self, path: Path) -> Dict[str, Any]:
        """Analyze PDF content using text extraction and OCR fallback for scanned documents"""
        # Extract text using the dedicated method
        pdf_text = self._extract_pdf_text(path)

        info = {}

        if len(pdf_text.strip()) > 100:
            self._analyze_pdf_text(pdf_text, info)
        else:
            self._analyze_scanned_pdf(path, info)

        return info

    def _analyze_pdf_text(self, pdf_text: str, info: Dict[str, Any]) -> None:
        """Analyze extracted PDF text content using AI"""
        if not self.ollama_analyzer:
            raise RuntimeError("Ollama analyzer not available for PDF text analysis")

        content_analysis = self.ollama_analyzer.analyze_text(
            pdf_text[:1000], "document"
        )
        info.update(content_analysis)

    def _analyze_scanned_pdf(self, path: Path, info: Dict[str, Any]) -> None:
        """Analyze scanned PDF using OCR extraction"""
        if not self.ocr_analyzer:
            raise RuntimeError("OCR analyzer not available for scanned PDF analysis")

        ocr_result = self.ocr_analyzer.extract_text_from_pdf(path)

        if not ocr_result["text"] or len(ocr_result["text"]) <= 50:
            raise RuntimeError("OCR failed to extract sufficient text from PDF")

        info["ocr_extracted"] = True
        info["ocr_confidence"] = ocr_result["confidence"]
        info["ocr_engine"] = ocr_result["engine"]

        if not self.ollama_analyzer:
            raise RuntimeError("Ollama analyzer not available for OCR text analysis")

        content_analysis = self.ollama_analyzer.analyze_text(
            ocr_result["text"][:1000], "document"
        )
        info.update(content_analysis)

    def _analyze_text_content(self, path: Path) -> Dict[str, Any]:
        """Analyze text file content"""
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read(2000)  # Read first 2000 chars

                if len(content.strip()) == 0:
                    return {"content": "empty_text_file"}

                # Basic content analysis
                lines = content.split("\n")
                info = {
                    "length": "short"
                    if len(content) < 500
                    else "medium"
                    if len(content) < 5000
                    else "long",
                    "lines": len(lines),
                }

                # AI content analysis
                if self.ollama_analyzer and len(content.strip()) > 50:
                    ai_analysis = self.ollama_analyzer.analyze_text(content, "document")
                    info.update(ai_analysis)

                return info
        except (OSError, UnicodeDecodeError, ValueError):
            return {"content": "text_analysis_failed"}

    def _analyze_csv_content(self, path: Path) -> Dict[str, Any]:
        """Analyze CSV file structure and content"""
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read(1500)  # Read more for better analysis

            if len(content.strip()) == 0:
                return {"content": "empty_csv"}

            lines = content.split("\n")
            columns = len(lines[0].split(",")) if lines else 0

            # AI content analysis for CSV
            info = {
                "columns": columns,
                "has_header": "," in lines[0] if lines else False,
            }

            if self.ollama_analyzer and len(content.strip()) > 50:
                ai_analysis = self.ollama_analyzer.analyze_text(content, "csv")
                info.update(ai_analysis)
            else:
                info["content"] = "csv_data"

            return info

        except (OSError, UnicodeDecodeError, ValueError):
            return {"content": "csv_analysis_failed"}

    def _analyze_code_content(self, path: Path) -> Dict[str, Any]:
        """Analyze code file content"""
        suffix = path.suffix.lower()
        language_map = {
            ".py": "python",
            ".js": "javascript",
            ".sh": "shell_script",
            ".html": "html",
            ".css": "css",
        }

        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read(1000)  # Read first 1000 chars

            info = {
                "language": language_map.get(suffix, "unknown"),
                "type": "script" if suffix == ".sh" else "source_code",
            }

            # AI analysis for scripts to understand what they do
            if self.ollama_analyzer and len(content.strip()) > 20:
                ai_analysis = self.ollama_analyzer.analyze_text(content, "script")
                info.update(ai_analysis)
            else:
                info["content"] = "code"

            return info
        except (OSError, UnicodeDecodeError, ValueError):
            return {
                "content": "code",
                "language": language_map.get(suffix, "unknown"),
                "type": "script" if suffix == ".sh" else "source_code",
            }

    def _analyze_word_content(self, path: Path) -> Dict[str, Any]:
        """Analyze Word document content"""
        try:
            doc = Document(path)

            # Count content
            paragraph_count = len(doc.paragraphs)
            word_count = sum(
                len(p.text.split()) for p in doc.paragraphs if p.text.strip()
            )

            info = {
                "content": "word_document",
                "paragraphs": paragraph_count,
                "words": word_count,
            }

            # Get a sample of text for AI analysis
            if self.ollama_analyzer and word_count > 0:
                sample_text = ""
                for p in doc.paragraphs[:5]:  # First 5 paragraphs
                    if p.text.strip():
                        sample_text += p.text + " "

                if len(sample_text) > 100:
                    ai_analysis = self.ollama_analyzer.analyze_text(
                        sample_text[:1000], "document"
                    )
                    info.update(ai_analysis)

            return info
        except (OSError, ImportError, ValueError):
            return {"content": "word_analysis_failed"}
